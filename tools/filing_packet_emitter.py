"""filing_packet_emitter.py — StegVerse Patents tool (v1)

Assembles a Patent-Center-ready PROVISIONAL filing packet for one
invention id. This is a TOOL invoked by the existing dispatcher or run
manually in CI — it is not a workflow (Core-Lite constraint).

The packet is everything a human needs to complete the boundary
crossing at https://patentcenter.uspto.gov in one sitting:

  filing_packets/<inv_id>/
    specification.docx        single DOCX: spec + CLAIMS + ABSTRACT
                              (Patent Center single-DOCX section format)
    cover_sheet_data.json     SB/16 provisional cover sheet field values
    fee_estimate.json         provisional fee by entity status (verify at
                              filing time against current USPTO schedule)
    FILING_CHECKLIST.md       ordered human steps, incl. writing the
                              application number back into the ledger
    PACKET_MANIFEST.json      sha256 receipt for every artifact + source
                              refs; status: ready-for-human-filing

Boundary doctrine: this tool NEVER contacts the USPTO. The submission
act is a human-gated transition. Post-filing, the human records the
application number in deadlines/deadlines.json, which arms downstream
ODP file-wrapper polling.

Usage:
  python tools/filing_packet_emitter.py --invention-id <id> \
      [--entity-status micro|small|large] [--root .]

Exit 0 on success; nonzero with a reason if inputs are missing.
"""

from __future__ import annotations

import argparse, datetime as dt, hashlib, json, pathlib, re, sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")

SIG = "filing-packet:v1"
PROVISIONAL_FEE_USD = {"large": 325, "small": 130, "micro": 65}


def sha256_of(path: pathlib.Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_markdown_sections(md: str):
    sections, current, buf = [], None, []
    for line in md.splitlines():
        m = re.match(r"^#{1,6}\s+(.*)$", line)
        if m:
            if buf or current is not None:
                sections.append((current, [p for p in "\n".join(buf).split("\n\n") if p.strip()]))
            current, buf = m.group(1).strip(), []
        else:
            buf.append(line)
    sections.append((current, [p for p in "\n".join(buf).split("\n\n") if p.strip()]))
    return sections


def extract_metadata(md: str) -> dict:
    meta = {}
    for m in re.finditer(r"\*\*([^*]+):\*\*\s*(.+?)\s*$", md, re.MULTILINE):
        meta[m.group(1).strip().lower()] = m.group(2).strip()
    return meta


def find_claims(root: pathlib.Path, inv_id: str, provisional_md: str):
    for cand in [root / "claims" / f"{inv_id}_claims.md",
                 root / "provisionals" / f"{inv_id}_claims.md"]:
        if cand.exists():
            txt = cand.read_text(encoding="utf-8")
            claims = re.findall(r"^\s*\d+\.\s+(.+?)(?=^\s*\d+\.|\Z)", txt,
                                re.MULTILINE | re.DOTALL)
            if claims:
                return [re.sub(r"\s+", " ", c).strip() for c in claims], str(cand.relative_to(root))
    for heading, paras in parse_markdown_sections(provisional_md):
        if heading and heading.lower().startswith("claims") and paras:
            body = "\n\n".join(paras)
            claims = re.findall(r"^\s*\d+\.\s+(.+?)(?=^\s*\d+\.|\Z)", body,
                                re.MULTILINE | re.DOTALL)
            if claims:
                return [re.sub(r"\s+", " ", c).strip() for c in claims], "provisional Claims section"
    return [], None


def build_specification_docx(out_path: pathlib.Path, title: str, inventors: str,
                             provisional_md: str, claims: list, warnings: list):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def body(text):
        doc.add_paragraph(re.sub(r"\s+", " ", text).strip())

    heading("Title of the Invention")
    body(title)
    skip = {"claims", "claims (placeholder)"}
    for h, paras in parse_markdown_sections(provisional_md):
        if h is None or h.lower() in skip or h.lower().startswith("provisional patent draft"):
            continue
        heading(h)
        for p in paras:
            body(p)

    heading("Claims")
    if claims:
        body("What is claimed is:")
        for i, c in enumerate(claims, 1):
            doc.add_paragraph(f"{i}. {c}")
    else:
        body("What is claimed is:")
        doc.add_paragraph("1. [PLACEHOLDER — claims not yet drafted. A provisional does not require claims, but at least one broad claim is recommended.]")
        warnings.append("claims-missing: placeholder claim inserted")

    heading("Abstract")
    summary = ""
    for h, paras in parse_markdown_sections(provisional_md):
        if h and h.lower() == "summary" and paras:
            summary = re.sub(r"\s+", " ", paras[0]).strip()
            break
    if summary and not summary.lower().startswith("high-level summary"):
        words = summary.split()
        if len(words) > 150:
            summary = " ".join(words[:150])
            warnings.append("abstract-truncated: summary exceeded 150 words")
        body(summary)
    else:
        body("[PLACEHOLDER — abstract of 150 words or fewer required. Populate the Summary section of the provisional draft.]")
        warnings.append("abstract-placeholder: Summary section is unpopulated template text")

    doc.save(str(out_path))


CHECKLIST = """# Filing Checklist — {inv_id}

Human boundary crossing. Estimated time: 10–15 minutes.

1. **Review** `specification.docx`. Resolve every PLACEHOLDER before filing.
   Warnings recorded in PACKET_MANIFEST.json: {warning_count}.
2. **Drawings**: if the spec references FIG. 1–N, prepare them as PDF and
   have the file ready. Provisionals accept informal drawings.
3. Sign in at https://patentcenter.uspto.gov (USPTO.gov account, ID.me
   verified). Select **New submission → Provisional (SB/16 equivalent)**.
4. Upload `specification.docx` as a single DOCX. Patent Center auto-detects
   the CLAIMS and ABSTRACT sections. Attach drawings PDF if any.
5. Enter cover sheet data from `cover_sheet_data.json` (title, inventors,
   correspondence, entity status).
6. **Entity status**: {entity_status}. If micro, complete the SB/15A
   certification in Patent Center. Verify eligibility before certifying.
7. **Fee**: verify the current provisional fee against the live USPTO fee
   schedule (estimate in `fee_estimate.json` is ${fee} as of the Jan 2025
   schedule) and pay.
8. Save the electronic filing receipt PDF into this packet directory as
   `uspto_filing_receipt.pdf`.
9. **Record the application number** in `deadlines/deadlines.json`:
   set `application_number`, `provisional_filed_utc` (actual filing date),
   `nonprovisional_due_utc` = filing date + 365 days, `pct_due_utc` = same,
   `status` = "provisional-filed".
10. Commit the updated packet directory and ledger. The commit is the
    receipt that the boundary crossing completed.
"""


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--invention-id", required=True)
    ap.add_argument("--entity-status", choices=["micro", "small", "large"], default="micro")
    ap.add_argument("--root", default=".")
    args = ap.parse_args()

    root = pathlib.Path(args.root).resolve()
    inv_id = args.invention_id
    warnings: list = []

    disclosure_path = root / "disclosures" / f"{inv_id}.md"
    provisional_path = root / "provisionals" / f"{inv_id}_provisional.md"
    if not provisional_path.exists():
        sys.exit(f"FAIL-CLOSED: no provisional draft at {provisional_path}")
    if not disclosure_path.exists():
        warnings.append(f"disclosure-missing: {disclosure_path.name} not found")

    provisional_md = provisional_path.read_text(encoding="utf-8")
    meta = extract_metadata(provisional_md)
    title = meta.get("title", inv_id)
    inventors = meta.get("inventors", "Rigel Randolph et al.")
    claims, claims_source = find_claims(root, inv_id, provisional_md)

    out_dir = root / "filing_packets" / inv_id
    out_dir.mkdir(parents=True, exist_ok=True)
    spec_path = out_dir / "specification.docx"
    build_specification_docx(spec_path, title, inventors, provisional_md, claims, warnings)

    cover = {
        "sig": "sb16-cover-data:v1",
        "application_type": "provisional",
        "title_of_invention": title,
        "inventors": [{"name": inventors, "residence": "TO-CONFIRM", "citizenship": "TO-CONFIRM"}],
        "correspondence": {"name": "TO-CONFIRM", "email": "rigel@stegverse.org", "address": "TO-CONFIRM"},
        "entity_status": args.entity_status,
        "government_interest": False,
        "filed_by": "pro se inventor",
    }
    (out_dir / "cover_sheet_data.json").write_text(json.dumps(cover, indent=2), encoding="utf-8")
    if "TO-CONFIRM" in json.dumps(cover):
        warnings.append("cover-sheet: TO-CONFIRM fields must be completed before filing")

    fee = PROVISIONAL_FEE_USD[args.entity_status]
    (out_dir / "fee_estimate.json").write_text(json.dumps({
        "sig": "fee-estimate:v1",
        "entity_status": args.entity_status,
        "provisional_filing_fee_usd": fee,
        "schedule_as_of": "2025-01",
        "verify_at_filing": True,
    }, indent=2), encoding="utf-8")

    (out_dir / "FILING_CHECKLIST.md").write_text(CHECKLIST.format(
        inv_id=inv_id, entity_status=args.entity_status, fee=fee,
        warning_count=len(warnings)), encoding="utf-8")

    artifacts = []
    for p in sorted(out_dir.iterdir()):
        if p.name == "PACKET_MANIFEST.json" or not p.is_file():
            continue
        artifacts.append({"path": p.name, "sha256": sha256_of(p), "bytes": p.stat().st_size})

    manifest = {
        "sig": SIG,
        "invention_id": inv_id,
        "created_utc": dt.datetime.now(dt.timezone.utc).isoformat().replace("+00:00", "Z"),
        "entity_status": args.entity_status,
        "source_refs": {
            "provisional": str(provisional_path.relative_to(root)),
            "disclosure": str(disclosure_path.relative_to(root)) if disclosure_path.exists() else None,
            "claims": claims_source,
        },
        "artifacts": artifacts,
        "warnings": warnings,
        "status": "ready-for-human-filing" if not warnings else "ready-with-warnings",
    }
    (out_dir / "PACKET_MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    ledger = root / "deadlines" / "deadlines.json"
    if ledger.exists():
        data = json.loads(ledger.read_text())
        for item in data.get("items", []):
            if item.get("invention_id") == inv_id and item.get("status") == "drafting":
                item["status"] = "packet-emitted"
                item["packet_emitted_utc"] = manifest["created_utc"]
        ledger.write_text(json.dumps(data, indent=2))

    print(f"[EMITTER] packet {manifest['status']} at {out_dir} ({len(artifacts)} artifacts, {len(warnings)} warnings)")
    for w in warnings:
        print(f"[EMITTER] WARN {w}")


if __name__ == "__main__":
    main()
